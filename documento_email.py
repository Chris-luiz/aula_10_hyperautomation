from docx import Document
import smtplib
from email.message import EmailMessage
from dotenv import load_dotenv
import os
from pathlib import Path

# Localiza o arquivo .env na pasta do projeto
BASE_DIR = Path(__file__).resolve().parent
load_dotenv(BASE_DIR / ".env")

def criar_documento():
    documento = Document()
    documento.add_heading("PORTAL FAKE SOLUÇÕES DIGITAIS", level=1)
    documento.add_heading("FICHA DE CADASTRO", level=2)
    documento.add_paragraph()
    documento.add_paragraph("1. Nome:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("2. Sobrenome:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("3. CPF:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("4. E-mail:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("5. Telefone:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("6. Data de Nascimento:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph("7. Endereço:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph()
    documento.add_paragraph("Assinatura:")
    documento.add_paragraph("____________________________________________________________")
    documento.add_paragraph()
    documento.add_paragraph("Data:")
    documento.add_paragraph("______/______/________")
    arquivo = "Ficha_Cadastro_Portal_Fake.docx" 
    documento.save(arquivo)
    return arquivo

def enviar_email(email_cliente, arquivo):
    # Lê as credenciais do .env

    remetente = os.getenv("EMAIL_REMETENTE")
    senha = os.getenv("EMAIL_SENHA")
# Verifica se encontrou as variáveis
    if not remetente:
        raise Exception("EMAIL_REMETENTE não encontrado no arquivo .env")
    if not senha:
        raise Exception("EMAIL_SENHA não encontrada no arquivo .env")

    print(f"Remetente: {remetente}")
    print(f"Senha encontrada: {len(senha)} caracteres")
    mensagem = EmailMessage()
    mensagem["Subject"] = "Ficha de Cadastro - Portal Fake Soluções Digitais" 
    mensagem["From"] = remetente
    mensagem["To"] = email_cliente
    mensagem.set_content( """Prezado(a), Segue em anexo a ficha de cadastro da Empresa Portal Fake Soluções Digitais. Solicitamos que todos os campos sejam preenchidos e que a ficha seja devolvida
    juntamente com: • Documento oficial com foto; • Comprovante de residência atualizado. Em caso de dúvidas, estamos à disposição. Atenciosamente, Portal Fake Soluções Digitais """)
    with open(arquivo, "rb") as documento:
        mensagem.add_attachment(
    documento.read(), maintype="application", subtype="vnd.openxmlformats-officedocument.wordprocessingml.document", filename=arquivo)
    servidor = smtplib.SMTP("smtp.gmail.com", 587)

    servidor.starttls()
    servidor.login(remetente, senha)
    servidor.send_message(mensagem)
    servidor.quit()
    print("E-mail enviado com sucesso!")

if __name__ == "__main__":
    arquivo = criar_documento()
    enviar_email( "vieirakarol302@gmail.com", arquivo)