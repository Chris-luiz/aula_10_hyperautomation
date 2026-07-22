from pathlib import Path
from playwright.sync_api import sync_playwright
from docx import Document
import random
import string
import time


def gerar_dados_aleatorios():
    nomes = ["Ana", "Bruno", "Carla", "Diego", "Eduarda", "Felipe", "Giovana", "Hugo"]
    sobrenomes = ["Silva", "Souza", "Oliveira", "Santos", "Pereira", "Costa", "Almeida"]
    ruas = ["Rua das Flores", "Av. Brasil", "Rua XV de Novembro", "Rua dos Andradas"]
    cidades = ["Manaus", "São Paulo", "Belo Horizonte", "Curitiba", "Recife"]

    nome = random.choice(nomes)
    sobrenome = random.choice(sobrenomes)

    def gerar_cpf():
        return "".join(random.choices(string.digits, k=3)) + "." + \
               "".join(random.choices(string.digits, k=3)) + "." + \
               "".join(random.choices(string.digits, k=3)) + "-" + \
               "".join(random.choices(string.digits, k=2))

    def gerar_telefone():
        return f"({random.randint(11,99)}) 9{random.randint(1000,9999)}-{random.randint(1000,9999)}"

    def gerar_nascimento():
        dia = random.randint(1, 28)
        mes = random.randint(1, 12)
        ano = random.randint(1960, 2005)
        return f"{ano}-{mes:02d}-{dia:02d}"  # formato usado em <input type="date">

    email = f"{nome.lower()}.{sobrenome.lower()}{random.randint(1,999)}@exemplo.com"
    endereco = f"{random.choice(ruas)}, {random.randint(10,999)} - {random.choice(cidades)}"

    return {
        "nome": nome,
        "sobrenome": sobrenome,
        "cpf": gerar_cpf(),
        "email": email,
        "telefone": gerar_telefone(),
        "nascimento": gerar_nascimento(),
        "endereco": endereco,
    }


def extrair_dados():
    dados_gerados = gerar_dados_aleatorios()

    with sync_playwright() as p:
        navigator = p.chromium.launch(headless=False)
        page = navigator.new_page()

        file = Path("../portal_fake/index.html").resolve()
        page.goto(f"file://{file}")
        page.click('#btnNovo')

        page.wait_for_timeout(1000)

        # Preenche o formulário com os dados aleatórios
        page.fill('#f_nome', dados_gerados["nome"])
        page.fill('#f_sobrenome', dados_gerados["sobrenome"])
        page.fill('#f_cpf', dados_gerados["cpf"])
        page.fill('#f_email', dados_gerados["email"])
        page.fill('#f_telefone', dados_gerados["telefone"])
        page.fill('#f_nascimento', dados_gerados["nascimento"])
        page.fill('#f_endereco', dados_gerados["endereco"])

        page.wait_for_timeout(1000)

        # Extrai os dados de volta do formulário
        dados = {
            "Nome":       page.locator("#f_nome").input_value(),
            "Sobrenome":  page.locator("#f_sobrenome").input_value(),
            "CPF":        page.locator("#f_cpf").input_value(),
            "E-mail":     page.locator("#f_email").input_value(),
            "Telefone":   page.locator("#f_telefone").input_value(),
            "Nascimento": page.locator("#f_nascimento").input_value(),
            "Endereço":   page.locator("#f_endereco").input_value(),
        }

        print("Dados Extraídos: ")
        for campo, valor in dados.items():
            print(f"{campo}: {valor}")

        navigator.close()

        gerar_docx(dados)


def gerar_docx(dados, caminho_saida="cadastro_extraido.docx"):
    doc = Document()
    doc.add_heading("Dados de Cadastro Extraídos", level=1)

    tabela = doc.add_table(rows=0, cols=2)
    tabela.style = "Light Grid Accent 1"

    for campo, valor in dados.items():
        linha = tabela.add_row().cells
        linha[0].text = campo
        linha[1].text = valor

    doc.save(caminho_saida)
    print(f"\nArquivo salvo em: {Path(caminho_saida).resolve()}")


if __name__ == '__main__':
    extrair_dados()